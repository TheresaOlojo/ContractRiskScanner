import os
import docx
import pandas as pd
import pdfplumber
import fitz  # PyMuPDF for PDF highlights
import openai
import matplotlib.pyplot as plt
from collections import Counter
from docx.enum.text import WD_COLOR_INDEX

# ==========================
#       OPENAI SETUP
# ==========================
openai.api_key = "YOUR_API_KEY_HERE"  # <- Replace with your key

# ==========================
#       RISK KEYWORDS
# ==========================
RISK_RULES = {
    "Indemnity": ["indemnify", "hold harmless", "liability", "compensate"],
    "Jurisdiction Risk": ["jurisdiction", "governing law", "venue", "court"],
    "Payment Risk": ["penalty", "late fee", "interest", "non-payment"],
    "Termination Risk": ["terminate", "termination", "breach"],
    "Confidentiality Risk": ["confidential", "non-disclosure", "nda"],
}

# ==================================
#        HELPER FUNCTIONS
# ==================================

def read_document(file_path):
    """Read DOCX or PDF and return text."""
    if file_path.lower().endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_path.lower().endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    else:
        print("Unsupported file format.")
        return ""

def ai_summarize_risk(clause):
    """Use OpenAI to summarize clause and suggest safer wording."""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful legal assistant."},
                {"role": "user", "content": f"Summarize this clause in plain English and suggest safer wording: '{clause}'"}
            ],
            temperature=0.5
        )
        summary = response.choices[0].message.content
        return summary
    except Exception as e:
        return f"AI summary error: {e}"

def scan_contract(text):
    """Scan text for keywords and return results with AI summary."""
    results = []
    lines = text.split(".")
    for i, line in enumerate(lines):
        sentence = line.strip().lower()
        for risk_label, keywords in RISK_RULES.items():
            for kw in keywords:
                if kw.lower() in sentence:
                    summary = ai_summarize_risk(line.strip())
                    results.append({
                        "Risk Type": risk_label,
                        "Keyword": kw,
                        "Sentence": line.strip(),
                        "Line Number": i + 1,
                        "AI Summary": summary
                    })
    return results

# ==================================
#        REPORT EXPORT FUNCTIONS
# ==================================

def save_csv_report(results, filename):
    os.makedirs("../reports", exist_ok=True)
    df = pd.DataFrame(results)
    output_path = f"../reports/{filename}_report.csv"
    df.to_csv(output_path, index=False)
    print(f"CSV report saved to: {output_path}")

def save_docx_report(results, original_text, filename):
    os.makedirs("../reports", exist_ok=True)
    doc = docx.Document()
    doc.add_heading(f"Contract Risk Report: {filename}", level=0)

    # Add summary
    num_risks = len(results)
    if num_risks == 0:
        score = "Green"
    elif num_risks <= 5:
        score = "Yellow"
    else:
        score = "Red"

    doc.add_paragraph(f"Number of risks found: {num_risks}")
    doc.add_paragraph(f"Overall Risk Score: {score}")

    doc.add_heading("Risks Found with AI Summary:", level=1)
    for r in results:
        p = doc.add_paragraph()
        p.add_run(f"Risk Type: {r['Risk Type']}\n")
        p.add_run(f"Keyword: {r['Keyword']}\n")
        p.add_run(f"Sentence: {r['Sentence']}\n")
        p.add_run(f"Line Number: {r['Line Number']}\n")
        p.add_run(f"AI Summary / Safer Wording:\n{r['AI Summary']}\n")
        p.add_run("-" * 50)

    # Add full contract text with highlights
    doc.add_page_break()
    doc.add_heading("Full Contract Text (Keywords Highlighted):", level=1)
    text_paragraphs = original_text.split("\n")
    for para in text_paragraphs:
        p = doc.add_paragraph()
        words = para.split(" ")
        for word in words:
            run = p.add_run(word + " ")
            for kw_list in RISK_RULES.values():
                for kw in kw_list:
                    if kw.lower() in word.lower():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    output_path = f"../reports/{filename}_report.docx"
    doc.save(output_path)
    print(f"DOCX report saved to: {output_path}")

def highlight_pdf(file_path, results, filename):
    """Highlight risky keywords in PDFs using PyMuPDF."""
    doc = fitz.open(file_path)
    keywords = list({r["Keyword"] for r in results})  # unique keywords
    for page in doc:
        for kw in keywords:
            text_instances = page.search_for(kw)
            for inst in text_instances:
                page.add_highlight_annot(inst)
    os.makedirs("../reports", exist_ok=True)
    output_path = f"../reports/{filename}_highlighted.pdf"
    doc.save(output_path)
    print(f"PDF with highlights saved to: {output_path}")

# ==================================
#        RISK DASHBOARD
# ==================================

def generate_dashboard(all_results):
    """Generate a bar chart of risk type distribution."""
    if not all_results:
        print("No results to generate dashboard.")
        return
    risk_types = [r["Risk Type"] for r in all_results]
    counts = Counter(risk_types)
    
    plt.figure(figsize=(8,5))
    plt.bar(counts.keys(), counts.values(), color='skyblue')
    plt.title("Risk Type Distribution")
    plt.ylabel("Count")
    plt.xlabel("Risk Type")
    plt.tight_layout()
    os.makedirs("../reports", exist_ok=True)
    plt.savefig("../reports/risk_dashboard.png")
    plt.show()
    print("Dashboard saved to ../reports/risk_dashboard.png")

# ==================================
#        RISK RULE MANAGEMENT
# ==================================

def show_rules():
    print("\n=== CURRENT RISK RULES ===")
    for category, keywords in RISK_RULES.items():
        print(f"\n{category}:")
        for kw in keywords:
            print(f"  - {kw}")
    print("\n")

def add_keyword():
    category = input("Enter risk category (or type NEW to create a new one): ")
    if category.lower() == "new":
        new_cat = input("Enter new category name: ")
        RISK_RULES[new_cat] = []
        category = new_cat
    keyword = input("Enter keyword to add: ").lower()
    if category not in RISK_RULES:
        print("Category does not exist.")
        return
    RISK_RULES[category].append(keyword)
    print(f"Keyword '{keyword}' added to '{category}'.")

def remove_keyword():
    category = input("Enter category to edit: ")
    if category not in RISK_RULES:
        print("Category does not exist.")
        return
    print(f"Keywords in {category}: {RISK_RULES[category]}")
    keyword = input("Enter keyword to remove: ").lower()
    if keyword in RISK_RULES[category]:
        RISK_RULES[category].remove(keyword)
        print("Keyword removed.")
    else:
        print("Keyword not found.")

# ==================================
#        SCAN FUNCTION
# ==================================

def scan_file():
    files = os.listdir("../contracts")
    files = [f for f in files if f.lower().endswith((".docx", ".pdf"))]

    if not files:
        print("No DOCX or PDF files found in ../contracts.")
        return

    print("\nAvailable contract files:")
    for i, f in enumerate(files):
        print(f"{i + 1}. {f}")

    choice = int(input("\nEnter file number to scan: ")) - 1
    file_name = files[choice]
    file_path = f"../contracts/{file_name}"
    base_filename = file_name.rsplit(".", 1)[0]

    print(f"\nReading file: {file_name}...\n")
    text = read_document(file_path)

    if not text.strip():
        print("No text extracted from file. Cannot scan.")
        return

    print("Scanning for risks with AI summaries...\n")
    results = scan_contract(text)

    if not results:
        print("No risks found. Contract appears clean.")
    else:
        print(f"{len(results)} risks found. Generating reports...")

    save_csv_report(results, base_filename)
    save_docx_report(results, text, base_filename)

    if file_path.lower().endswith(".pdf"):
        highlight_pdf(file_path, results, base_filename)

    # Generate dashboard for this single file
    generate_dashboard(results)

# ==================================
#            MAIN MENU
# ==================================

def main_menu():
    while True:
        print("""
=== CONTRACT RISK SCANNER MENU ===

1. Scan a Contract
2. View Risk Rules
3. Add a Keyword
4. Remove a Keyword
5. Exit
""")
        choice = input("Enter choice: ")

        if choice == "1":
            scan_file()
        elif choice == "2":
            show_rules()
        elif choice == "3":
            add_keyword()
        elif choice == "4":
            remove_keyword()
        elif choice == "5":
            print("Goodbye!")
            break
        else:
            print("Invalid choice. Try again.")

# ==================================
#             RUN
# ==================================

if __name__ == "__main__":
    main_menu()
